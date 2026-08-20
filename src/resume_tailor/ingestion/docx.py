from __future__ import annotations

import hashlib
from pathlib import Path

# DOCX is an optional extra.  Keep this import inside the adapter so a base
# installation can still start the CLI, inspect diagnostics, and ingest the
# formats it supports.
from resume_tailor.errors import ValidationError
from resume_tailor.ingestion.safety import safe_docx_members
from resume_tailor.models.cv import CVBullet, CVDocument, CVEntry, CVSection, SourceLocation
from resume_tailor.models.ingestion import (
    IngestionResult,
    IngestionStatistics,
    IngestionWarning,
    SourceFormat,
)

_CANONICAL = {
    "contact",
    "profile",
    "links",
    "summary",
    "education",
    "experience",
    "projects",
    "skills",
    "certifications",
    "achievements",
}


def _name(value: str) -> str:
    return " ".join(value.lower().split())


def ingest_docx(path: Path) -> IngestionResult:
    safe_docx_members(path)
    raw = path.read_bytes()
    try:
        from docx import Document

        doc = Document(str(path))
    except ModuleNotFoundError as exc:
        raise ValidationError(
            "DOCX support is not installed; reinstall with 'resume-tailor[docx]'"
        ) from exc
    except Exception as exc:
        raise ValidationError("could not parse DOCX safely") from exc
    source = str(path)
    sections = []
    warnings = []
    current = None
    paragraph_count = 0
    table_count = len(doc.tables)
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        paragraph_count += 1
        loc = SourceLocation(file=source, format="docx", part="body", paragraph=paragraph_count)
        style = (p.style.name if p.style is not None else "").lower()
        canonical = _name(text)
        if canonical in _CANONICAL or style.startswith("heading"):
            section_name = canonical if canonical in _CANONICAL else "unknown"
            if section_name == "unknown":
                warnings.append(
                    IngestionWarning(
                        code="unknown_heading", message=f"Unrecognized heading: {text}"
                    )
                )
            current = CVSection(
                name=section_name,
                entries=[CVEntry(heading=text, source_location=loc)],
                source_location=loc,
                level=1,
            )
            sections.append(current)
            continue
        if current is None:
            current = CVSection(
                name="unknown",
                entries=[CVEntry(heading="Uncategorized", source_location=loc)],
                source_location=loc,
                level=1,
            )
            sections.append(current)
        entry = current.entries[0]
        bullet = "list" in style or (p._p.pPr is not None and p._p.pPr.numPr is not None)
        if bullet:
            entry.bullets.append(CVBullet(text=text, source_location=loc))
        else:
            entry.text.append(text)
    for table_i, table in enumerate(doc.tables, 1):
        for row_i, row in enumerate(table.rows, 1):
            for cell_i, cell in enumerate(row.cells, 1):
                text = cell.text.strip()
                if text:
                    loc = SourceLocation(
                        file=source,
                        format="docx",
                        locator=f"body/t:{table_i}/r:{row_i}/c:{cell_i}",
                        label=f"table {table_i}, row {row_i}, cell {cell_i}",
                        part="body",
                    )
                    if current is None:
                        current = CVSection(
                            name="unknown",
                            entries=[CVEntry(heading="Uncategorized", source_location=loc)],
                            source_location=loc,
                            level=1,
                        )
                        sections.append(current)
                    current.entries[0].text.append(text)
    if not sections:
        raise ValidationError("DOCX contains no extractable resume content")
    document = CVDocument(
        source_file=source, sections=sections, warnings=[w.message for w in warnings]
    )
    return IngestionResult(
        source_format=SourceFormat.DOCX,
        document=document,
        warnings=warnings,
        statistics=IngestionStatistics(
            paragraphs=paragraph_count,
            tables=table_count,
            characters=sum(len(p.text) for p in doc.paragraphs),
        ),
        source_hash=f"sha256:{hashlib.sha256(raw).hexdigest()}",
    )
